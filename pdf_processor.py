"""
PDF Document Processing Pipeline - Phase 1
Multimodal Alignment of Noisy Image-Text Pairs using Weak Supervision

This module implements the core PDF processing pipeline including:
- PDF document loading and iteration
- Image extraction using PyMuPDF
- Text chunking and preprocessing
- Caption extraction and linking
"""

import os
import json
import re
import fitz  # PyMuPDF
import pdfplumber
import spacy
from pathlib import Path
from typing import List, Dict, Any, Tuple, Optional
from tqdm import tqdm
import numpy as np
from PIL import Image
from dotenv import load_dotenv
from docx import Document
import tempfile
import subprocess
from collections import Counter

try:
    import pytesseract  # type: ignore
except Exception:  # pragma: no cover
    pytesseract = None

# Load environment variables
load_dotenv()


class PDFProcessor:
    """Main class for processing PDF documents and extracting multimodal data."""

    def __init__(self, input_dir: str, output_dir: str):
        """
        Initialize the PDF processor.

        Args:
            input_dir: Directory containing PDF files
            output_dir: Directory to save processed data
        """
        self.input_dir = Path(input_dir)
        self.output_dir = Path(output_dir)
        self.images_dir = self.output_dir / "images"

        # Load configuration from environment variables
        self.max_file_size_mb = int(os.getenv("MAX_FILE_SIZE_MB", "50"))
        self.allowed_file_types = os.getenv("ALLOWED_FILE_TYPES", "pdf,docx,doc").split(
            ","
        )
        self.api_key = os.getenv("OPENAI_API_KEY")
        self.log_level = os.getenv("LOG_LEVEL", "INFO")
        self.language = os.getenv("LANGUAGE", "nl")  # Dutch by default
        self.use_ocr_fallback = os.getenv("USE_OCR_FALLBACK", "False").lower() == "true"

        # Create output directories
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.images_dir.mkdir(parents=True, exist_ok=True)

        # Global counters for unique IDs
        self.global_image_counter = 0
        self.global_chunk_counter = 0

        # Initialize spaCy model based on language
        self.nlp = self._load_spacy_model()

        # Storage for extracted data
        self.image_metadata = []
        self.text_chunks = []
        self.lexical_components = set()

        # Section tracking for structural hierarchy
        self.current_section = None
        self.current_level = None
        self.section_stats = {"sections": 0, "subsections": 0}

    def detect_section_level(self, text: str) -> Tuple[Optional[str], Optional[int]]:
        """Detect section headers like 'Section 2.1 Maintenance' or '2.1.3 Rotor Assembly'."""
        # Pattern to match hierarchical section numbering
        pattern = re.compile(
            r"^(Section|Hoofdstuk|Hfdst\.|§)?\s*(\d+(?:\.\d+)+)\s*[:\-]?\s*(.*)$",
            re.IGNORECASE,
        )
        match = pattern.match(text.strip())
        if match:
            title = match.group(0).strip()
            # Count how many numbers in the hierarchy to define level
            nums = re.findall(r"\d+", match.group(2))
            level = len(nums)
            return title, level
        return None, None

    def filter_invalid_bboxes(self, images):
        """Remove or ignore images that have zero bounding boxes."""
        if not images:
            return images
        valid = [
            img
            for img in images
            if img.get("bbox") and any((c or 0) != 0 for c in img["bbox"])
        ]
        invalid_count = len(images) - len(valid)
        if invalid_count > 0:
            total = len(images)
            ratio = (invalid_count / total) if total else 0.0
            print(
                f"Filtered out {invalid_count} images with zero bounding boxes ({ratio:.2%})."
            )
        return valid

    def _render_page_to_image(self, page, dpi: int = 200):
        """Render a PDF page to a PIL Image."""
        try:
            pix = page.get_pixmap(dpi=dpi)
            mode = "RGB" if pix.alpha == 0 else "RGBA"
            img = Image.frombytes(mode, (pix.width, pix.height), pix.samples)
            return img
        except Exception:
            return None

    def extract_ocr_bboxes(self, page) -> list:
        """Extract OCR text boxes from a page using pytesseract.image_to_data.

        Returns a list of dicts: {"text": str, "bbox": [x0,y0,x1,y1]}
        Coordinates are in rendered image space; we map them back to PDF coords.
        """
        if not self.use_ocr_fallback or pytesseract is None:
            return []

        img = self._render_page_to_image(page, dpi=200)
        if img is None:
            return []

        try:
            data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
        except Exception:
            return []

        # Map image coords back to PDF coords via scale factors
        page_rect = page.rect
        scale_x = page_rect.width / float(img.width)
        scale_y = page_rect.height / float(img.height)

        results = []
        n = len(data.get("text", []))
        for i in range(n):
            txt = (data["text"][i] or "").strip()
            if not txt:
                continue
            try:
                x = int(data["left"][i])
                y = int(data["top"][i])
                w = int(data["width"][i])
                h = int(data["height"][i])
            except Exception:
                continue
            # Tesseract y grows downward; PDF origin is top-left in fitz coordinate system for page.rect mapping
            x0 = x * scale_x
            y0 = y * scale_y
            x1 = (x + w) * scale_x
            y1 = (y + h) * scale_y
            results.append({"text": txt, "bbox": [x0, y0, x1, y1]})

        return results

    def _load_spacy_model(self):
        """Load appropriate spaCy model based on language setting."""
        model_map = {
            "en": "en_core_web_sm",
            "nl": "nl_core_news_sm",  # Dutch
            "de": "de_core_news_sm",  # German
            "fr": "fr_core_news_sm",  # French
        }

        model_name = model_map.get(self.language, "en_core_web_sm")

        try:
            nlp = spacy.load(model_name)
            print(f"✓ Loaded spaCy model: {model_name}")
            return nlp
        except OSError:
            print(f"Warning: {model_name} not found. Trying to download...")
            try:
                import subprocess

                subprocess.run(
                    [os.sys.executable, "-m", "spacy", "download", model_name],
                    check=True,
                )
                nlp = spacy.load(model_name)
                print(f"✓ Downloaded and loaded: {model_name}")
                return nlp
            except Exception as e:
                print(f"Failed to load {model_name}: {e}")
                print("Falling back to basic text processing (no NLP)")
                return None

    def _convert_word_to_pdf(self, word_path: Path) -> Optional[Path]:
        """Convert a Word document to PDF locally using LibreOffice."""
        try:
            tmp_dir = Path(tempfile.mkdtemp(prefix="word2pdf_"))
            pdf_out = tmp_dir / f"{word_path.stem}.pdf"

            # Use LibreOffice
            try:
                soffice_path = os.getenv(
                    "SOFFICE_PATH",
                    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
                )
                subprocess.run(
                    [
                        soffice_path,
                        "--headless",
                        "--convert-to",
                        "pdf",
                        str(word_path),
                        "--outdir",
                        str(tmp_dir),
                    ],
                    check=True,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    timeout=int(os.getenv("WORD_TO_PDF_TIMEOUT_SECONDS", "120")),
                )
                if pdf_out.exists() and pdf_out.stat().st_size > 0:
                    return pdf_out
            except subprocess.TimeoutExpired:
                print(f"LibreOffice conversion timed out for {word_path}")
            except FileNotFoundError:
                print(
                    "LibreOffice (soffice) not found. Install it or set SOFFICE_PATH to its binary."
                )
            except Exception as e:
                print(f"LibreOffice conversion failed for {word_path}: {e}")

        except Exception as e:
            print(f"Word->PDF conversion failed for {word_path}: {e}")

        return None

    def _log_image_summary(self, manual_id: str) -> None:
        native = sum(
            1
            for m in self.image_metadata
            if m.get("manual_id") == manual_id and m.get("bbox_source") == "native"
        )
        vector = sum(
            1
            for m in self.image_metadata
            if m.get("manual_id") == manual_id and m.get("bbox_source") == "vector"
        )
        total = sum(1 for m in self.image_metadata if m.get("manual_id") == manual_id)
        print(f"Image extraction summary for {manual_id}:")
        print(f"  Native images: {native}")
        print(f"  Vector figures: {vector}")
        print(f"  Total entries: {total}")

    def _log_section_summary(self, manual_id: str) -> None:
        """Log section hierarchy statistics for a document."""
        # Get total pages processed for this manual
        doc_chunks = [c for c in self.text_chunks if c.get("manual_id") == manual_id]
        total_pages = len(set(c.get("page") for c in doc_chunks if c.get("page")))

        print(
            f"Structure extraction for {manual_id}: "
            f"Detected {self.section_stats['sections']} sections and "
            f"{self.section_stats['subsections']} subsections across "
            f"{total_pages} pages."
        )

    def process_all_documents(self) -> None:
        """Process all supported documents in the input directory."""
        # Reset previous results to avoid duplication
        self.image_metadata = []
        self.text_chunks = []
        self.lexical_components = set()
        self.global_image_counter = 0
        self.global_chunk_counter = 0

        for file in [
            "image_metadata.json",
            "text_chunks.json",
            "lexical_components.json",
        ]:
            path = self.output_dir / file
            if path.exists():
                path.unlink()

        # Remove all images in the images directory
        for image in self.images_dir.glob("*"):
            try:
                image.unlink()
            except Exception as e:
                print(f"Error removing image {image}: {e}")
                continue

        # Find all supported file types
        all_files = []
        for file_type in self.allowed_file_types:
            pattern = f"*.{file_type}"
            files = list(self.input_dir.glob(pattern))
            all_files.extend(files)

        if not all_files:
            print(f"No supported files found in {self.input_dir}")
            print(f"Supported types: {', '.join(self.allowed_file_types)}")
            return

        print(f"Found {len(all_files)} files to process")

        for file_path in tqdm(all_files, desc="Processing documents"):
            try:
                self.process_single_document(file_path)
            except Exception as e:
                print(f"Error processing {file_path}: {e}")
                continue

        # Save all extracted data
        self.save_extracted_data()

    def process_single_document(self, file_path: Path) -> None:
        """Process a single document (PDF or Word)."""
        manual_id = file_path.stem
        file_extension = file_path.suffix.lower()

        print(f"Processing {manual_id} ({file_extension})...")

        # Reset section tracking for new document
        self.current_section = None
        self.current_level = None
        self.section_stats = {"sections": 0, "subsections": 0}

        if file_extension == ".pdf":
            # Extract images
            self.extract_images_from_pdf(file_path, manual_id)
            # Extract text chunks and captions
            self.extract_text_chunks_from_pdf(file_path, manual_id)
            self._log_image_summary(manual_id)
            self._log_section_summary(manual_id)
        elif file_extension in [".docx", ".doc"]:
            # Convert to PDF to obtain proper bounding boxes, then process as PDF
            converted_pdf = self._convert_word_to_pdf(file_path)
            if converted_pdf is not None and converted_pdf.exists():
                # Use the PDF pipeline (images + text with bbox + captions)
                self.extract_images_from_pdf(converted_pdf, manual_id)
                self.extract_text_chunks_from_pdf(converted_pdf, manual_id)
                self._log_image_summary(manual_id)
                self._log_section_summary(manual_id)
            else:
                # Do not proceed without bbox-capable pipeline
                raise RuntimeError(
                    "Word->PDF conversion failed; aborting to avoid zero-bbox Word extraction."
                )
        else:
            print(f"Unsupported file type: {file_extension}")

    def process_single_pdf(self, pdf_path: Path) -> None:
        """Process a single PDF file (legacy method)."""
        self.process_single_document(pdf_path)

    def extract_images_from_pdf(self, pdf_path: Path, manual_id: str) -> None:
        """Extract images from PDF using PyMuPDF."""
        try:
            doc = fitz.open(pdf_path)

            for page_num in range(len(doc)):
                page = doc[page_num]
                images = page.get_images(full=True)

                for img_idx, img in enumerate(images):
                    try:
                        # Get image XREF
                        xref = img[0]

                        # Extract image
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        image_ext = base_image["ext"]

                        # Get image bounding box
                        # Primary: use MuPDF-provided image rects
                        rects = page.get_image_rects(xref)
                        img_rect = rects[0] if rects else None
                        bbox_source = None
                        bbox = (
                            [img_rect.x0, img_rect.y0, img_rect.x1, img_rect.y1]
                            if img_rect
                            else None
                        )
                        if bbox is not None:
                            bbox_source = "native"

                        # Fallback: parse page text dict to find image block with matching xref
                        if bbox is None:
                            try:
                                raw = page.get_text("dict")
                                fallback_bbox = None
                                for block in raw.get("blocks", []):
                                    if block.get("type") == 1:  # image block
                                        info = block.get("image", {})
                                        block_xref = info.get("xref")
                                        if block_xref == xref and "bbox" in block:
                                            fb = block["bbox"]
                                            fallback_bbox = [fb[0], fb[1], fb[2], fb[3]]
                                            break
                                if fallback_bbox is not None:
                                    bbox = fallback_bbox
                                    bbox_source = "dict_fallback"
                                else:
                                    # As a last resort, use first image block bbox on the page
                                    img_blocks = [
                                        b
                                        for b in raw.get("blocks", [])
                                        if b.get("type") == 1 and "bbox" in b
                                    ]
                                    if img_blocks:
                                        fb = img_blocks[
                                            min(img_idx, len(img_blocks) - 1)
                                        ]["bbox"]
                                        bbox = [fb[0], fb[1], fb[2], fb[3]]
                                        bbox_source = "dict_fallback"
                            except Exception:
                                bbox = None

                        if bbox is None:
                            # OCR-based approximation: find nearest caption OCR box, then search nearby OCR regions
                            ocr_boxes = self.extract_ocr_bboxes(page)
                            approx = None
                            # Try to find caption-like OCR text
                            caption_regex = re.compile(
                                r"\b(Fig\.|Figure|Figuur|Afb\.|Afbeelding|Foto)\b",
                                re.IGNORECASE,
                            )
                            caption_boxes = [
                                b
                                for b in ocr_boxes
                                if caption_regex.search(b["text"]) is not None
                            ]
                            # If a caption box exists, use a region above/below as proxy
                            if caption_boxes:
                                cb = caption_boxes[0]["bbox"]
                                # Heuristic: image likely above the caption; expand region above
                                y_top = max(0, cb[1] - (page.rect.height * 0.35))
                                approx = [
                                    max(0, cb[0] - 50),
                                    y_top,
                                    min(page.rect.width, cb[2] + 50),
                                    cb[1] - 5,
                                ]
                            else:
                                # As last resort, use union of all OCR boxes on page which are not tiny
                                big = [
                                    b["bbox"]
                                    for b in ocr_boxes
                                    if (b["bbox"][2] - b["bbox"][0])
                                    * (b["bbox"][3] - b["bbox"][1])
                                    > 2000
                                ]
                                if big:
                                    x0 = min(bb[0] for bb in big)
                                    y0 = min(bb[1] for bb in big)
                                    x1 = max(bb[2] for bb in big)
                                    y1 = max(bb[3] for bb in big)
                                    approx = [x0, y0, x1, y1]
                            bbox = [0, 0, 0, 0]
                            bbox_source = "unknown"

                        # Create unique filename
                        image_filename = (
                            f"{manual_id}_p{page_num + 1}_img{img_idx}.{image_ext}"
                        )
                        image_path = self.images_dir / image_filename

                        # Save image
                        with open(image_path, "wb") as img_file:
                            img_file.write(image_bytes)

                        # Store metadata with section context
                        image_metadata = {
                            "image_id": f"{manual_id}_p{page_num + 1}_img{img_idx}",
                            "manual_id": manual_id,
                            "page": page_num + 1,
                            "bbox": bbox,
                            "bbox_source": bbox_source or "unknown",
                            "caption": None,  # Will be filled later
                            "filename": image_filename,
                            "image_type": "raster_image",
                            "section_title": self.current_section,
                            "section_level": self.current_level,
                        }

                        self.image_metadata.append(image_metadata)
                        self.global_image_counter += 1

                    except Exception as e:
                        print(
                            f"Error extracting image {img_idx} from page {page_num}: {e}"
                        )
                        continue

                # Vector drawing detection
                try:
                    drawings = page.get_drawings()
                    v_idx = 0
                    for d in drawings:
                        rect = d.get("rect")
                        if not rect:
                            continue
                        w = float(rect.x1 - rect.x0)
                        h = float(rect.y1 - rect.y0)
                        if w < 5 or h < 5:
                            continue
                        vector_meta = {
                            "image_id": f"{manual_id}_p{page_num + 1}_vector{v_idx}",
                            "manual_id": manual_id,
                            "page": page_num + 1,
                            "bbox": [rect.x0, rect.y0, rect.x1, rect.y1],
                            "bbox_source": "vector",
                            "caption": None,
                            "filename": None,
                            "image_type": "vector_figure",
                            "section_title": self.current_section,
                            "section_level": self.current_level,
                        }
                        self.image_metadata.append(vector_meta)
                        self.global_image_counter += 1
                        v_idx += 1
                except Exception:
                    pass

            doc.close()

        except Exception as e:
            print(f"Error processing PDF {pdf_path}: {e}")

    def extract_images_from_word(self, word_path: Path, manual_id: str) -> None:
        """Extract images from Word document."""
        try:
            doc = Document(word_path)

            # Extract images from the document
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        # Get image data
                        image_data = rel.target_part.blob

                        # Determine image format
                        if image_data.startswith(b"\xff\xd8\xff"):
                            ext = "jpg"
                        elif image_data.startswith(b"\x89PNG"):
                            ext = "png"
                        elif image_data.startswith(b"GIF"):
                            ext = "gif"
                        elif image_data.startswith(b"BM"):
                            ext = "bmp"
                        else:
                            ext = "png"  # Default fallback

                        # Create unique filename
                        image_filename = (
                            f"{manual_id}_img{self.global_image_counter}.{ext}"
                        )
                        image_path = self.images_dir / image_filename

                        # Save image
                        with open(image_path, "wb") as img_file:
                            img_file.write(image_data)

                        # Get image dimensions
                        try:
                            from PIL import Image
                            from io import BytesIO

                            img_pil = Image.open(BytesIO(image_data))
                            width, height = img_pil.size
                        except:
                            width, height = 0, 0

                        # Store metadata
                        image_metadata = {
                            "image_id": f"{manual_id}_img{self.global_image_counter}",
                            "manual_id": manual_id,
                            "page": 1,  # Word docs don't have clear page breaks
                            "bbox": [0, 0, 0, 0],  # No bbox available for Word images
                            "bbox_source": "unknown",
                            "dimensions": {"width": width, "height": height},
                            "image_order": self.global_image_counter,
                            "caption": None,  # Will be filled later
                            "filename": image_filename,
                            "image_type": "raster_image",
                        }

                        self.image_metadata.append(image_metadata)
                        self.global_image_counter += 1

                    except Exception as e:
                        print(f"Error extracting image from Word document: {e}")
                        continue

        except Exception as e:
            print(f"Error extracting images from Word document {word_path}: {e}")

    def extract_text_chunks_from_pdf(self, pdf_path: Path, manual_id: str) -> None:
        """Extract text chunks and captions using pdfplumber."""
        try:
            with pdfplumber.open(pdf_path) as pdf:
                total_pages = len(pdf.pages)
                for page_num, page in enumerate(pdf.pages):
                    # Extract text blocks with section tracking
                    text_blocks = self.extract_text_blocks(
                        page, manual_id, page_num + 1, total_pages
                    )
                    self.text_chunks.extend(text_blocks)

                    # Extract and link captions
                    self.extract_and_link_captions(page, manual_id, page_num + 1)

        except Exception as e:
            print(f"Error extracting text from {pdf_path}: {e}")

    def extract_text_chunks_from_word(self, word_path: Path, manual_id: str) -> None:
        """Extract text chunks from Word document."""
        try:
            doc = Document(word_path)

            # Process each paragraph
            for para_idx, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    # Check if it's an instruction line
                    if self.is_instruction_line(text):
                        chunk_metadata = {
                            "chunk_id": f"{manual_id}_p1_c{para_idx}",
                            "manual_id": manual_id,
                            "page": 1,  # Word docs don't have clear page breaks
                            "bbox": [0, 0, 0, 0],  # No bbox available for Word
                            "text": text,
                        }
                        self.text_chunks.append(chunk_metadata)
                        self.global_chunk_counter += 1
                    else:
                        # Split by sentences for non-instruction text
                        sentences = self.split_by_sentences(text)
                        for sent_idx, sentence in enumerate(sentences):
                            if sentence.strip():
                                chunk_metadata = {
                                    "chunk_id": f"{manual_id}_p1_c{para_idx}_{sent_idx}",
                                    "manual_id": manual_id,
                                    "page": 1,
                                    "bbox": [0, 0, 0, 0],
                                    "text": sentence.strip(),
                                }
                                self.text_chunks.append(chunk_metadata)
                                self.global_chunk_counter += 1

            # Process tables
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    row_text = " ".join(
                        [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    )
                    if row_text:
                        chunk_metadata = {
                            "chunk_id": f"{manual_id}_table{table_idx}_row{row_idx}",
                            "manual_id": manual_id,
                            "page": 1,
                            "bbox": [0, 0, 0, 0],
                            "text": row_text,
                        }
                        self.text_chunks.append(chunk_metadata)
                        self.global_chunk_counter += 1

            # Extract and link captions for Word documents
            self.extract_and_link_captions_word(doc, manual_id)

        except Exception as e:
            print(f"Error extracting text from Word document {word_path}: {e}")

    def extract_text_blocks(
        self, page, manual_id: str, page_num: int, total_pages: int
    ) -> List[Dict[str, Any]]:
        """Extract text blocks and split into instruction-level chunks."""
        text_blocks = []

        try:
            # Get text objects with bounding boxes
            # Try to get words with reasonable tolerances to maximize detection
            try:
                words = page.extract_words(
                    x_tolerance=2, y_tolerance=2, keep_blank_chars=False
                )
            except Exception:
                words = page.extract_words()

            if not words:
                # Fallback: extract plain text and split by lines
                plain_text = page.extract_text()
                added_any = False
                if plain_text:
                    lines = plain_text.split("\n")
                    for line_idx, line in enumerate(lines):
                        if line.strip():
                            # Try OCR to infer bbox of this line
                            ocr_boxes = self.extract_ocr_bboxes(page)
                            bbox = [0, 0, 0, 0]
                            if ocr_boxes:
                                # pick the closest OCR box text-wise (simple containment match)
                                candidates = [
                                    b
                                    for b in ocr_boxes
                                    if line.strip() in b["text"]
                                    or b["text"] in line.strip()
                                ]
                                if candidates:
                                    bbox = candidates[0]["bbox"]
                            chunk_metadata = {
                                "chunk_id": f"{manual_id}_p{page_num}_c{line_idx}",
                                "manual_id": manual_id,
                                "page": page_num,
                                "bbox": bbox,
                                "text": line.strip(),
                            }
                            text_blocks.append(chunk_metadata)
                            self.global_chunk_counter += 1
                            added_any = True
                return text_blocks

            # Group words into lines based on vertical proximity
            lines = self.group_words_into_lines(words)

            # Split lines into instruction-level chunks
            chunks = self.split_into_instruction_chunks(lines)

            for chunk_idx, chunk_text in enumerate(chunks):
                if chunk_text.strip():
                    # Calculate bounding box for the chunk
                    # Heuristic: use the bbox of the source line if sentence-level mapping is uncertain
                    # Find the first line that contributed text to this chunk
                    bbox = [0, 0, 0, 0]
                    for line in lines:
                        line_text = " ".join([w["text"] for w in line]).strip()
                        if not line_text:
                            continue
                        if (
                            chunk_text.strip() in line_text
                            or line_text in chunk_text.strip()
                        ):
                            bbox = self.calculate_chunk_bbox(line)
                            break
                    # Fallback: union of all words whose text overlaps token-wise with the chunk
                    if bbox == [0, 0, 0, 0]:
                        chunk_tokens = set(
                            t for t in re.split(r"\s+", chunk_text.strip()) if t
                        )
                        chunk_words = [
                            w
                            for line in lines
                            for w in line
                            if w.get("text") and w["text"] in chunk_tokens
                        ]
                        if chunk_words:
                            bbox = self.calculate_chunk_bbox(chunk_words)

                    # Detect if this chunk is a section header
                    section_title, section_level = self.detect_section_level(chunk_text)
                    if section_title:
                        self.current_section = section_title
                        self.current_level = section_level
                        if section_level == 1:
                            self.section_stats["sections"] += 1
                        else:
                            self.section_stats["subsections"] += 1

                    chunk_metadata = {
                        "chunk_id": f"{manual_id}_p{page_num}_c{chunk_idx}",
                        "manual_id": manual_id,
                        "page": page_num,
                        "bbox": bbox,
                        "text": chunk_text.strip(),
                        "section_title": self.current_section,
                        "section_level": self.current_level,
                    }

                    text_blocks.append(chunk_metadata)
                    self.global_chunk_counter += 1

        except Exception as e:
            print(f"Error extracting text blocks from page {page_num}: {e}")
            # Fallback to simple text extraction
            try:
                plain_text = page.extract_text()
                if plain_text:
                    lines = plain_text.split("\n")
                    for line_idx, line in enumerate(lines):
                        if line.strip():
                            chunk_metadata = {
                                "chunk_id": f"{manual_id}_p{page_num}_c{line_idx}",
                                "manual_id": manual_id,
                                "page": page_num,
                                "bbox": [0, 0, 0, 0],
                                "text": line.strip(),
                            }
                            text_blocks.append(chunk_metadata)
                            self.global_chunk_counter += 1
            except Exception as e2:
                print(f"Fallback text extraction also failed: {e2}")

        return text_blocks

    def group_words_into_lines(self, words: List[Dict]) -> List[List[Dict]]:
        """Group words into lines based on vertical proximity."""
        if not words:
            return []

        # Sort words by vertical position
        words.sort(key=lambda w: w.get("top", 0))

        lines = []
        current_line = [words[0]]
        line_height = words[0].get("bottom", 0) - words[0].get("top", 0)
        tolerance = line_height * 0.5

        for word in words[1:]:
            # Check if word is on the same line
            if abs(word.get("top", 0) - current_line[0].get("top", 0)) <= tolerance:
                current_line.append(word)
            else:
                # Sort current line by horizontal position
                current_line.sort(key=lambda w: w.get("x0", w.get("left", 0)))
                lines.append(current_line)
                current_line = [word]

        # Add the last line
        if current_line:
            current_line.sort(key=lambda w: w.get("x0", w.get("left", 0)))
            lines.append(current_line)

        return lines

    def split_into_instruction_chunks(self, lines: List[List[Dict]]) -> List[str]:
        """Split lines into instruction-level chunks."""
        chunks = []

        for line in lines:
            line_text = " ".join([word["text"] for word in line])

            # Check for instruction patterns
            if self.is_instruction_line(line_text):
                chunks.append(line_text)
            else:
                # For non-instruction lines, try to split by sentences
                sentences = self.split_by_sentences(line_text)
                chunks.extend(sentences)

        return chunks

    def is_instruction_line(self, text: str) -> bool:
        """Check if a line represents an instruction."""
        # Dutch instruction patterns
        dutch_patterns = [
            r"^\d+\.",  # Numbered list
            r"^[•·▪▫]",  # Bullet points
            r"^[a-zA-Z]\.",  # Lettered list
            r"^(Stap|Procedure|Instructie|Opmerking|Waarschuwing|Voorzichtigheid|Let op|Controleer|Verwijder|Installeer|Vervang|Controle|Onderhoud)",
            r"^[A-Z][a-z]+:",  # Bold headers
        ]

        # English patterns (fallback)
        english_patterns = [
            r"^\d+\.",  # Numbered list
            r"^[•·▪▫]",  # Bullet points
            r"^[a-zA-Z]\.",  # Lettered list
            r"^(Step|Procedure|Instruction|Note|Warning|Caution|Check|Remove|Install|Replace|Maintenance)",
            r"^[A-Z][a-z]+:",  # Bold headers
        ]

        patterns = dutch_patterns if self.language == "nl" else english_patterns

        for pattern in patterns:
            if re.match(pattern, text.strip(), re.IGNORECASE):
                return True

        return False

    def split_by_sentences(self, text: str) -> List[str]:
        """Split text into sentences."""
        if not self.nlp:
            # Simple sentence splitting if spaCy not available
            sentences = re.split(r"[.!?]+", text)
            return [s.strip() for s in sentences if s.strip()]

        doc = self.nlp(text)
        return [sent.text.strip() for sent in doc.sents if sent.text.strip()]

    def calculate_chunk_bbox(self, words: List[Dict]) -> List[float]:
        """Calculate bounding box for a chunk of words."""
        if not words:
            return [0, 0, 0, 0]

        # pdfplumber words use x0,x1,top,bottom; support both
        left = min(word.get("x0", word.get("left", 0)) for word in words)
        top = min(word.get("top", 0) for word in words)
        right = max(word.get("x1", word.get("right", 0)) for word in words)
        bottom = max(word.get("bottom", 0) for word in words)

        return [left, top, right, bottom]

    def extract_and_link_captions(self, page, manual_id: str, page_num: int) -> None:
        """Extract figure captions and link them to images."""
        # Get all text on the page
        page_text = page.extract_text()

        if not page_text:
            return

        # Find caption patterns (Dutch and English)
        dutch_caption_patterns = [
            r"Fig\.?\s*\d+[:\s]+.*?(?=\n|$)",
            r"Figuur\s*\d+[:\s]+.*?(?=\n|$)",
            r"Afb\.?\s*\d+[:\s]+.*?(?=\n|$)",
            r"Afbeelding\s*\d+[:\s]+.*?(?=\n|$)",
            r"Foto\s*\d+[:\s]+.*?(?=\n|$)",
        ]

        english_caption_patterns = [
            r"Fig\.?\s*\d+[:\s]+.*?(?=\n|$)",
            r"Figure\s*\d+[:\s]+.*?(?=\n|$)",
            r"Abb\.?\s*\d+[:\s]+.*?(?=\n|$)",
            r"Image\s*\d+[:\s]+.*?(?=\n|$)",
        ]

        caption_patterns = (
            dutch_caption_patterns
            if self.language == "nl"
            else english_caption_patterns
        )

        captions = []
        for pattern in caption_patterns:
            matches = re.finditer(pattern, page_text, re.IGNORECASE | re.MULTILINE)
            for match in matches:
                caption_text = match.group().strip()
                captions.append(
                    {"text": caption_text, "start": match.start(), "end": match.end()}
                )

        # Link captions to nearest images on the same page
        page_images = [
            img
            for img in self.image_metadata
            if img["manual_id"] == manual_id and img["page"] == page_num
        ]

        for caption in captions:
            # Find the closest image (simplified - could be improved with better spatial analysis)
            if page_images:
                # For now, assign to the first image on the page
                # In a more sophisticated implementation, you'd analyze spatial relationships
                page_images[0]["caption"] = caption["text"]

    def extract_and_link_captions_word(self, doc, manual_id: str) -> None:
        """Extract figure captions from Word document and link them to images."""
        # Get all text from the document
        all_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                all_text.append(paragraph.text.strip())

        page_text = "\n".join(all_text)

        if not page_text:
            return

        # Find caption patterns (Dutch and English)
        dutch_caption_patterns = [
            r"Fig\.?\s*\d+[:\s]+.*?(?=\n|$)",
            r"Figuur\s*\d+[:\s]+.*?(?=\n|$)",
            r"Afb\.?\s*\d+[:\s]+.*?(?=\n|$)",
            r"Afbeelding\s*\d+[:\s]+.*?(?=\n|$)",
            r"Foto\s*\d+[:\s]+.*?(?=\n|$)",
        ]

        english_caption_patterns = [
            r"Fig\.?\s*\d+[:\s]+.*?(?=\n|$)",
            r"Figure\s*\d+[:\s]+.*?(?=\n|$)",
            r"Abb\.?\s*\d+[:\s]+.*?(?=\n|$)",
            r"Image\s*\d+[:\s]+.*?(?=\n|$)",
        ]

        caption_patterns = (
            dutch_caption_patterns
            if self.language == "nl"
            else english_caption_patterns
        )

        captions = []
        for pattern in caption_patterns:
            matches = re.finditer(pattern, page_text, re.IGNORECASE | re.MULTILINE)
            for match in matches:
                caption_text = match.group().strip()
                captions.append(
                    {"text": caption_text, "start": match.start(), "end": match.end()}
                )

        # Link captions to images from the same document
        doc_images = [
            img for img in self.image_metadata if img["manual_id"] == manual_id
        ]

        for caption in captions:
            # Find the closest image based on caption number
            if doc_images:
                # Try to match caption number with image order
                caption_match = re.search(r"(\d+)", caption["text"])
                if caption_match:
                    caption_num = int(caption_match.group(1))
                    # Find image with matching order (caption_num - 1 for 0-based indexing)
                    target_image_idx = caption_num - 1
                    if 0 <= target_image_idx < len(doc_images):
                        doc_images[target_image_idx]["caption"] = caption["text"]
                    else:
                        # Fallback to first image if number doesn't match
                        doc_images[0]["caption"] = caption["text"]
                else:
                    # No number found, assign to first available image
                    doc_images[0]["caption"] = caption["text"]

    def preprocess_text(self, text: str) -> str:
        """Preprocess text for lexical component extraction."""
        # Remove hyphenation across lines
        text = re.sub(r"-\s*\n\s*", "", text)

        # Normalize whitespace
        text = re.sub(r"\s+", " ", text)

        # Remove page numbers and headers/footers (basic patterns)
        text = re.sub(r"^\d+\s*$", "", text, flags=re.MULTILINE)

        return text.strip()

    def extract_lexical_components(self, text: str) -> List[str]:
        """Extract lexical components from text using spaCy - nouns only for visual entities."""
        if not self.nlp:
            return []

        doc = self.nlp(text)
        components = []

        for token in doc:
            # Extract only nouns (visual entities that can appear in images)
            # Apply stricter filtering to catch artifacts
            lemma_lower = token.lemma_.lower().strip()

            # Filter criteria:
            # - Must be a noun
            # - Not a stop word
            # - Not punctuation
            # - Minimum length of 4 characters (to filter artifacts like "pken", "proce", "visionplaa")
            # - Must be alphanumeric (no special characters except hyphens for compound words)
            # - Must have at least one letter (not just digits)
            if (
                token.pos_ == "NOUN"
                and not token.is_stop
                and not token.is_punct
                and len(lemma_lower) >= 4
                and (
                    lemma_lower.replace("-", "").replace("_", "").isalnum()
                    or "-" in lemma_lower
                )
                and any(c.isalpha() for c in lemma_lower)  # At least one letter
            ):
                components.append(lemma_lower)

        return components

    def save_extracted_data(self) -> None:
        """Save all extracted data to JSON files."""
        # Step 1: Filter images with zero bboxes before saving
        self.image_metadata = self.filter_invalid_bboxes(self.image_metadata)

        # Save image metadata
        with open(self.output_dir / "image_metadata.json", "w") as f:
            json.dump(self.image_metadata, f, indent=2)

        # Save text chunks
        with open(self.output_dir / "text_chunks.json", "w") as f:
            json.dump(self.text_chunks, f, indent=2)

        # Extract and save lexical components
        all_text = " ".join([chunk["text"] for chunk in self.text_chunks])
        processed_text = self.preprocess_text(all_text)
        lexical_components = self.extract_lexical_components(processed_text)

        # Count frequencies and sort by frequency (descending), then alphabetically
        component_counts = Counter(lexical_components)

        # Create list of (term, count) tuples, sorted by frequency (desc), then alphabetically
        sorted_components = sorted(
            component_counts.items(),
            key=lambda x: (-x[1], x[0]),  # Sort by -count (desc), then by term (asc)
        )

        lexical_data = {
            "total_components": len(sorted_components),
            "total_occurrences": sum(component_counts.values()),
            "components": [
                {"term": term, "count": count} for term, count in sorted_components
            ],
        }

        with open(self.output_dir / "lexical_components.json", "w") as f:
            json.dump(lexical_data, f, indent=2)

        print(
            f"Saved {len(self.image_metadata)} images, {len(self.text_chunks)} text chunks, and {len(sorted_components)} unique lexical components (nouns only)"
        )


def main():
    """Main function to run the document processing pipeline."""
    input_dir = "data/raw/manuals"
    output_dir = "data/processed"

    processor = PDFProcessor(input_dir, output_dir)
    processor.process_all_documents()


if __name__ == "__main__":
    main()
